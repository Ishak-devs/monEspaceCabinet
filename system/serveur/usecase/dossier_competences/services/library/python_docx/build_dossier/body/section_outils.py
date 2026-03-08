from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt, Cm

def section_outils(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("OUTILS")
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = True

    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002060"/>')
    p._element.get_or_add_pPr().append(shd)

    for categorie in data.get('Logiciels_par_titre', []):
        for outil in categorie.get('logiciels_outils', []):
            p = doc.add_paragraph(outil, style='List Bullet')
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Pt(20)