from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt, Cm

COULEUR_PRINCIPALE = RGBColor(0x1B, 0x4A, 0x8A)
POLICE = "Arial"

def section_competences(doc, data):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Expériences")
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = True

    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="002060"/>')
    p._element.get_or_add_pPr().append(shd)

    for exp in data.get('Expériences', []):
        table = doc.add_table(rows=1, cols=3)
        table.columns[0].width = Cm(7)
        table.columns[1].width = Cm(4)
        table.columns[2].width = Cm(6)

        cell_gauche = table.cell(0, 0)
        p_gauche = cell_gauche.paragraphs[0]
        p_gauche.alignement = WD_ALIGN_PARAGRAPH.LEFT
        run = p_gauche.add_run(exp.get('Nom_Entreprise', ''))
        run.font.name = POLICE
        run.font.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = COULEUR_PRINCIPALE

        cell_milieu = table.cell(0, 1)
        p_milieu = cell_milieu.paragraphs[0]
        p_milieu.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p_milieu.add_run(exp.get('Durée_expérience', ''))
        run2.font.name = POLICE
        run2.font.size = Pt(10)
        run2.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

        cell_droite = table.cell(0, 2)
        p_droite = cell_droite.paragraphs[0]
        p_droite.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run3 = p_droite.add_run(exp.get('Dates', ''))
        run3.font.name = POLICE
        run3.font.size = Pt(10)
        run3.font.color.rgb = RGBColor(0x77, 0x77, 0x77)