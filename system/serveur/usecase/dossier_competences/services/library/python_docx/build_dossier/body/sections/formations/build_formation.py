from docx.shared import Pt

def build_formation(doc, data):
    p = doc.add_paragraph("FORMATIONS")
    p.style.font.bold = True

    for diplome in data.get('Diplômes', []):
        p = doc.add_paragraph()
        run = p.add_run(f"{diplome.get('Année')} : {diplome.get('Diplôme')}")
        run.bold = True
        p.add_run(f"\n{diplome.get('École')} - {diplome.get('Lieu')}")
        p.paragraph_format.space_after = Pt(8)