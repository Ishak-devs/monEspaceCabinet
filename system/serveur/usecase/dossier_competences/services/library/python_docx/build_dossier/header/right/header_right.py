from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.shared import Pt, RGBColor

POLICE = "Arial"
COULEUR_PRINCIPALE = RGBColor(0x1B, 0x4A, 0x8A)
COULEUR_GRIS = RGBColor(0x77, 0x77, 0x77)

def header_right(data, cell_droite):

    cell_droite.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p1 = cell_droite.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run1 = p1.add_run(data.get('nom', 'NOM PRENOM'))
    run1.font.name = POLICE
    run1.font.bold = True
    run1.font.size = Pt(14)
    run1.font.color.rgb = COULEUR_PRINCIPALE

    p2 = cell_droite.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(data.get('poste', 'Consultant IT'))
    run2.font.name = POLICE
    run2.font.size = Pt(11)

    p3 = cell_droite.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = p3.add_run(f"{data.get('annees_xp', '0')} ans d'expérience")
    run3.font.name = POLICE
    run3.font.size = Pt(10)
    run3.font.color.rgb = COULEUR_GRIS

    return header_right(data)