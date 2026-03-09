from docx.shared import Pt, RGBColor


def display_mission(doc, exp):
    p_mission = doc.add_paragraph()

    p_mission.add_run("MISSION : ").bold = True
    p_mission.add_run(str(exp.get('Description_Mission', '')))

    p_mission.paragraph_format.space_before = Pt(10)
    p_mission.paragraph_format.space_after = Pt(10)

    p_mission.paragraph_format.keep_with_next = True
