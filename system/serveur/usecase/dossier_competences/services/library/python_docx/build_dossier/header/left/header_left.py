from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm


def header_left(logo_path, cell_gauche):
    cell_gauche.margin_left = Cm(0)

    tcPr = cell_gauche._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:left w:w="0" w:type="dxa"/></tcMar>')

    tcPr.append(tcMar)
    paragraph = cell_gauche.paragraphs[0]

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.left_indent = Cm(0)

    run_logo = paragraph.add_run()
    run_logo.add_picture(logo_path, width=Cm(3))

    # delete_bord()