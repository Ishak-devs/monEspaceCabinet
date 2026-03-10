from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Cm

from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.formations.cells.left_cells import \
    left_cells
from usecase.dossier_competences.services.library.python_docx.build_dossier.body.sections.formations.cells.right_cells import \
    right_cells


def build_section_formation(doc, data):

    formations = data.get('Diplômes_Et_Formations') or []

    if not formations or not formations[0].get('Diplôme'):
        return

    p = doc.add_paragraph()

    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run("FORMATIONS")
    run.font.color.rgb = RGBColor(255, 255, 255)
    run.font.bold = True

    p._element.get_or_add_pPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="002060"/>'))
    p.paragraph_format.keep_with_next = True

    for diplome in data.get('Diplômes_Et_Formations', []):
        table = doc.add_table(rows=1, cols=2)

        table.style = None
        cells = table.rows[0].cells
        table.columns[0].width = Cm(14)
        table.columns[1].width = Cm(3)

        for row in table.rows:
            row._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")} w:val="1"/>'))

        if diplome.get('Diplôme'):
            left_cells(cells, diplome)

        if diplome.get('Année'):
            right_cells(cells, diplome)