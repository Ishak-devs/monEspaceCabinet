from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import RGBColor, Pt, Cm

from usecase.dossier_competences.services.library.python_docx.build_dossier.body.header_section.header_section import \
    header_section


def build_section_outils(doc, data):
    outils_data = data.get('Logiciels_Et_Outils_Sans_Indiquer_Le_Niveau', [])
    if outils_data and outils_data[0].get('Liste_Logiciels'):

        header_section(doc, "LOGICIELS ET OUTILS")

        for categorie in outils_data:
            titre = categorie.get('Catégorie', '')
            logiciels_outils = list(dict.fromkeys(categorie.get('Liste_Logiciels', [])))

            if titre and logiciels_outils:
                p_titre = doc.add_paragraph()
                run_t = p_titre.add_run(f"{titre.upper()} :")
                run_t.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
                run_t.bold = True
                run_t.underline = True
                p_titre.paragraph_format.space_after = Pt(2)

                for logiciel in logiciels_outils:
                    p_bullet = doc.add_paragraph(logiciel, style='List Bullet')
                    p_bullet.paragraph_format.space_after = Pt(2)
                    

                p_bullet.paragraph_format.space_after = Pt(6)
